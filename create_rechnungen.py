from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from copy import deepcopy
import subprocess

GREEN = RGBColor(0x2D, 0x4A, 0x3E)

def create_rechnung(filename, title_suffix, rows, gesamtbetrag, ust_note):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Logo
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.add_run().add_picture('logo.jpg', width=Cm(4))
    logo.paragraph_format.space_after = Pt(10)
    
    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f'RECHNUNG')
    run.font.size = Pt(16)
    run.font.color.rgb = GREEN
    run.bold = True
    t.paragraph_format.space_after = Pt(2)
    
    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(title_suffix)
    run.font.size = Pt(10)
    run.font.color.rgb = GREEN
    run.bold = True
    sub.paragraph_format.space_after = Pt(6)
    
    # Practice info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = info.add_run('Charlotte von Sichart · tiefenpsychologisch fundierte Psychotherapeutin\nKameruner Straße 43 · 13351 Berlin\nSteuer-Nr.: _____________________')
    run.font.size = Pt(9)
    info.paragraph_format.space_after = Pt(10)
    
    # Patient section
    ph = doc.add_paragraph()
    run = ph.add_run('Patient/in')
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN
    run.bold = True
    ph.paragraph_format.space_after = Pt(4)
    
    for field in ['Name, Vorname:', 'Anschrift:', 'Geburtsdatum:', 'Diagnose (ICD-10):']:
        p = doc.add_paragraph()
        run = p.add_run(f'{field} __________________________________________________')
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)
    
    # Rechnungsnummer
    p = doc.add_paragraph()
    run = p.add_run('Rechnungsnummer: ________________     Rechnungsdatum: ________________')
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    
    # Table
    table = doc.add_table(rows=len(rows)+2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Datum', 'GOÄ-Ziffer', 'Leistungsbeschreibung', 'Faktor', 'Betrag']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): '2D4A3E'
        })
        shading.append(shading_elem)
    
    for idx, (datum, ziffer, beschreibung, faktor, betrag) in enumerate(rows):
        row = table.rows[idx + 1]
        row.cells[0].text = datum
        row.cells[1].text = ziffer
        row.cells[2].text = beschreibung
        row.cells[3].text = faktor
        row.cells[4].text = betrag
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Gesamtbetrag row
    last_row = table.rows[-1]
    last_row.cells[3].text = ''
    run = last_row.cells[3].paragraphs[0].add_run('Gesamtbetrag:')
    run.bold = True
    run.font.size = Pt(9)
    last_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    last_row.cells[4].text = ''
    run = last_row.cells[4].paragraphs[0].add_run(gesamtbetrag)
    run.bold = True
    run.font.size = Pt(9)
    last_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(2)
        row.cells[2].width = Cm(7)
        row.cells[3].width = Cm(2)
        row.cells[4].width = Cm(2.5)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # Payment info
    zh = doc.add_paragraph()
    run = zh.add_run('Zahlungsinformationen')
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN
    run.bold = True
    zh.paragraph_format.space_after = Pt(4)
    
    p = doc.add_paragraph()
    run = p.add_run('Bankverbindung: IBAN _________________________________________')
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    
    p = doc.add_paragraph()
    run = p.add_run('Zahlungsfrist: 14 Tage ab Rechnungsdatum')
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    
    # USt + GOÄ note
    p = doc.add_paragraph()
    run = p.add_run('Diese Rechnung wurde gemäß der Gebührenordnung für Ärzte (GOÄ) erstellt.')
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    
    p = doc.add_paragraph()
    run = p.add_run(ust_note)
    run.font.size = Pt(9)
    run.italic = True
    p.paragraph_format.space_after = Pt(2)
    
    # Footer
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = fp.add_run('Charlotte von Sichart · tiefenpsychologisch fundierte Psychotherapeutin · Kameruner Straße 43 · 13351 Berlin · www.von-sichart.de')
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0xC9, 0xA8, 0x6C)
    
    doc.save(filename)
    print(f'Saved: {filename}')


# 1. GOP 3,5x MIT PPB
create_rechnung(
    'Rechnungsvorlage_3.5x_mit_PPB.docx',
    'GOP 3,5-facher Satz mit Psychopathologischem Befund',
    [
        ('________', '861', 'Tiefenpsychologisch fundierte Psychotherapie,\nEinzelbehandlung, mind. 50 Min.', '3,5x', '140,76 €'),
        ('________', '801', 'Psychiatrische Untersuchung /\nErhebung des psychischen Befundes', '3,5x', '81,11 €'),
    ],
    '221,87 €',
    'Kein Ausweis der Umsatzsteuer gemäß § 4 Nr. 14 UStG (Heilbehandlung).'
)

# 2. GOP 3,5x OHNE PPB
create_rechnung(
    'Rechnungsvorlage_3.5x_ohne_PPB.docx',
    'GOP 3,5-facher Satz ohne Psychopathologischen Befund',
    [
        ('________', '861', 'Tiefenpsychologisch fundierte Psychotherapie,\nEinzelbehandlung, mind. 50 Min.', '3,5x', '140,76 €'),
    ],
    '140,76 €',
    'Kein Ausweis der Umsatzsteuer gemäß § 4 Nr. 14 UStG (Heilbehandlung).'
)

# 3. GOP 2,3x MIT PPB (Kostenerstattung)
create_rechnung(
    'Rechnungsvorlage_2.3x_mit_PPB.docx',
    'GOP 2,3-facher Satz mit Psychopathologischem Befund (Kostenerstattung)',
    [
        ('________', '812a', 'Psychotherapeutische Kurzzeittherapie,\nEinzelbehandlung, mind. 50 Min.', '2,3x', '134,06 €'),
        ('________', '801a', 'Erhebung des aktuellen\npsychischen Befundes', '2,3x', '33,52 €'),
    ],
    '167,58 €',
    'Kein Ausweis der Umsatzsteuer gemäß § 4 Nr. 14 UStG (Heilbehandlung).'
)

# Convert all to PDF
for f in ['Rechnungsvorlage_3.5x_mit_PPB.docx', 'Rechnungsvorlage_3.5x_ohne_PPB.docx', 'Rechnungsvorlage_2.3x_mit_PPB.docx']:
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', f], check=True)
    print(f'PDF: {f.replace(".docx", ".pdf")}')
