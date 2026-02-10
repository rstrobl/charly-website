#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

GOLD = RGBColor(0xC9, 0xA8, 0x6C)
GREEN = RGBColor(0x2D, 0x4A, 0x3E)
GREY = RGBColor(0x66, 0x66, 0x66)
LOGO = os.path.join(os.path.dirname(__file__), 'logo.jpg')
FOOTER_TEXT = "Charlotte von Sichart · Systemische Paartherapeutin · Kameruner Straße 43 · 13351 Berlin · www.von-sichart.de"

def setup_doc(narrow_margins=True):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)
    
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    return doc

def add_logo(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(LOGO, width=Cm(4))
    p.paragraph_format.space_after = Pt(2)

def add_contact_line(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Kameruner Straße 43 · 13351 Berlin\nTel. 0163 436 93 69 · charlotte.vonsichart@outlook.de")
    r.font.size = Pt(8.5)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(2)

def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━" * 60)
    r.font.size = Pt(8)
    r.font.color.rgb = GOLD
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = GREEN
    p.paragraph_format.space_after = Pt(2)

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(6)

def add_heading(doc, text, size=Pt(11)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = size
    r.font.color.rgb = GREEN
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_text(doc, text, size=Pt(10), bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = size
    r.bold = bold
    p.paragraph_format.space_after = Pt(3)
    return p

def add_field(doc, label, underline_len=40):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.font.size = Pt(10)
    r2 = p.add_run("_" * underline_len)
    r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_checkbox(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"☐  {text}")
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_signature_line(doc, label):
    p = doc.add_paragraph()
    r = p.add_run("_" * 45)
    r.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(0)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(label)
    r2.font.size = Pt(8)
    r2.font.color.rgb = GREY
    p2.paragraph_format.space_after = Pt(2)

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # separator
    r = p.add_run("━" * 50 + "\n")
    r.font.size = Pt(7)
    r.font.color.rgb = GOLD
    r2 = p.add_run(FOOTER_TEXT)
    r2.font.size = Pt(7)
    r2.font.color.rgb = GREY

def add_header(doc):
    add_logo(doc)
    add_contact_line(doc)
    add_separator(doc)

# ===== 1. SCHWEIGEPFLICHTENTBINDUNG =====
def create_schweigepflicht():
    doc = setup_doc()
    add_header(doc)
    add_title(doc, "SCHWEIGEPFLICHTENTBINDUNG")
    add_subtitle(doc, "Entbindung von der Schweigepflicht gemäß § 203 StGB")
    
    add_heading(doc, "Angaben zur Person")
    add_field(doc, "Name, Vorname", 50)
    add_field(doc, "Geburtsdatum", 50)
    add_field(doc, "Anschrift", 50)
    
    add_heading(doc, "Entbindung der Schweigepflicht gegenüber")
    add_text(doc, "Ich entbinde Charlotte von Sichart, Systemische Paartherapeutin, von ihrer Schweigepflicht gegenüber folgenden Personen/Institutionen:")
    
    add_checkbox(doc, "Hausarzt/Facharzt: ___________________________________________")
    add_checkbox(doc, "Anschrift: ___________________________________________")
    add_checkbox(doc, "Krankenkasse: ___________________________________________")
    add_checkbox(doc, "Andere Therapeut*innen: ___________________________________________")
    add_checkbox(doc, "Angehörige: ___________________________________________")
    
    add_heading(doc, "Umfang der Entbindung")
    add_text(doc, "Die Entbindung umfasst die Weitergabe folgender Informationen:")
    add_checkbox(doc, "Befunde und Diagnosen")
    add_checkbox(doc, "Therapieverlauf und -fortschritt")
    add_checkbox(doc, "Therapieempfehlungen")
    add_checkbox(doc, "Sonstige: ___________________________________________")
    
    add_heading(doc, "Widerruf")
    add_text(doc, "Diese Schweigepflichtentbindung kann jederzeit schriftlich widerrufen werden. Der Widerruf gilt ab dem Zeitpunkt des Zugangs. Bereits erfolgte Übermittlungen bleiben vom Widerruf unberührt.", Pt(9))
    
    add_separator(doc)
    add_field(doc, "Ort, Datum", 50)
    add_signature_line(doc, "Unterschrift Patient/in")
    add_signature_line(doc, "Charlotte von Sichart, Therapeutin")
    
    add_footer(doc)
    path = os.path.join(os.path.dirname(__file__), 'Schweigepflichtentbindung_Charlotte_von_Sichart.docx')
    doc.save(path)
    return path

# ===== 2. RECHNUNGSVORLAGE =====
def create_rechnung():
    doc = setup_doc()
    add_header(doc)
    add_title(doc, "RECHNUNG")
    
    # Practice info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Charlotte von Sichart · Systemische Paartherapeutin\nKameruner Straße 43 · 13351 Berlin\nSteuernummer: _________________________")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(6)
    
    add_heading(doc, "Patient/in")
    add_field(doc, "Name, Vorname", 50)
    add_field(doc, "Anschrift", 50)
    add_field(doc, "Geburtsdatum", 50)
    
    # Rechnungsnummer etc
    p = doc.add_paragraph()
    r = p.add_run("Rechnungsnummer: ________________     Rechnungsdatum: ________________")
    r.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Table
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Datum', 'GOÄ-Ziffer', 'Leistungsbeschreibung', 'Faktor', 'Betrag']
    rows_data = [
        ['________', '812a', 'Tiefenpsychologische Behandlung (Einzelsitzung, 50 Min.)', '2,3x', '120,44 €'],
        ['________', '801a', 'Psychiatrische Untersuchung', '2,3x', '53,62 €'],
        ['________', '870a', 'Verhaltenstherapeutische Intervention', '2,3x', '92,50 €'],
    ]
    
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = GREEN
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            if col_idx == 4:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif col_idx in (0, 1, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Last row - total
    last = table.rows[4]
    for i in range(4):
        last.cells[i].text = ''
    p = last.cells[3].paragraphs[0]
    r = p.add_run("Gesamtbetrag:")
    r.bold = True
    r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = last.cells[4].paragraphs[0]
    p.text = ''
    r = p.add_run("266,56 €")
    r.bold = True
    r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(2.2)
        row.cells[1].width = Cm(2.2)
        row.cells[2].width = Cm(7)
        row.cells[3].width = Cm(1.5)
        row.cells[4].width = Cm(2.5)
    
    add_text(doc, "")
    add_heading(doc, "Zahlungsinformationen")
    add_text(doc, "Bankverbindung: IBAN _________________________________________")
    add_text(doc, "Zahlungsfrist: 14 Tage ab Rechnungsdatum")
    
    add_separator(doc)
    p = doc.add_paragraph()
    r = p.add_run("Diese Rechnung wurde gemäß der Gebührenordnung für Ärzte (GOÄ) erstellt. Bitte reichen Sie diese Rechnung bei Ihrer Krankenversicherung bzw. Beihilfestelle zur Erstattung ein.")
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    r.italic = True
    p.paragraph_format.space_after = Pt(2)
    
    add_footer(doc)
    path = os.path.join(os.path.dirname(__file__), 'Rechnungsvorlage_Charlotte_von_Sichart.docx')
    doc.save(path)
    return path

# ===== 3. THERAPIEVERTRAG PAARTHERAPIE =====
def create_paartherapie():
    doc = setup_doc()
    add_header(doc)
    add_title(doc, "THERAPIEVERTRAG · PAARTHERAPIE")
    add_subtitle(doc, "Systemische Paartherapie")
    
    add_text(doc, "Nach ausführlicher Information und Aufklärung über die Bedingungen einer Paartherapie wird zwischen")
    
    p = doc.add_paragraph()
    r = p.add_run("Charlotte von Sichart, Systemische Paartherapeutin")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(", nachfolgend Therapeutin genannt, und")
    r2.font.size = Pt(10)
    
    add_field(doc, "Partner/in 1: Name, Vorname", 40)
    add_field(doc, "Geburtsdatum", 40)
    add_field(doc, "Anschrift", 40)
    
    add_field(doc, "Partner/in 2: Name, Vorname", 40)
    add_field(doc, "Geburtsdatum", 40)
    add_field(doc, "Anschrift", 40)
    
    add_text(doc, "nachfolgend gemeinsam das Paar genannt, folgender Therapievertrag geschlossen:")
    
    add_heading(doc, "1. Gegenstand der Therapie")
    add_text(doc, "Gegenstand ist die Durchführung einer systemischen Paartherapie. Die Therapie dient der Bearbeitung partnerschaftlicher Themen, der Verbesserung der Kommunikation und der gemeinsamen Weiterentwicklung der Beziehung.")
    
    add_heading(doc, "2. Rahmenbedingungen")
    add_text(doc, "Die Sitzungen dauern in der Regel 90 Minuten. Die Häufigkeit wird individuell nach Vereinbarung festgelegt. Die Sitzungen finden in der Praxis, Kameruner Straße 43, 13351 Berlin, statt.")
    
    add_heading(doc, "3. Honorar")
    add_text(doc, "Das Honorar beträgt 220,00 € pro Sitzung (90 Minuten). Paartherapie ist eine Privatleistung und wird nicht von den gesetzlichen Krankenkassen übernommen. Das Honorar ist nach jeder Sitzung fällig. Beide Partner haften gesamtschuldnerisch.")
    
    add_heading(doc, "4. Absageregelung")
    add_text(doc, "Vereinbarte Termine sind spätestens 48 Stunden vorher abzusagen. Bei nicht rechtzeitiger Absage wird das volle Honorar in Rechnung gestellt. Dies gilt auch bei Absage durch nur einen Partner.")
    
    add_heading(doc, "5. Schweigepflicht")
    add_text(doc, "Die Therapeutin unterliegt der gesetzlichen Schweigepflicht (§ 203 StGB). Alle Informationen, die im Rahmen der Therapie mitgeteilt werden, werden streng vertraulich behandelt. Beide Partner erklären sich mit der gemeinsamen Therapie einverstanden.")
    
    add_heading(doc, "6. Besonderheit der Paartherapie")
    add_text(doc, "Die Therapeutin ist beiden Partnern gleichermaßen verpflichtet und versteht sich als allparteiliche Begleiterin. Grundsätzlich werden keine Geheimnisse zwischen den Partnern in der Therapie bewahrt. Informationen, die der Therapeutin von einem Partner allein mitgeteilt werden, können in den gemeinsamen Sitzungen thematisiert werden.")
    
    add_heading(doc, "7. Einzelgespräche")
    add_text(doc, "Einzelgespräche mit einem Partner sind nach gemeinsamer Absprache möglich. Die Inhalte können in den Paarsitzungen aufgegriffen werden (siehe Punkt 6).")
    
    add_heading(doc, "8. Beendigung der Therapie")
    add_text(doc, "Die Therapie kann jederzeit von allen Beteiligten beendet werden. Es wird empfohlen, eine Beendigung im Rahmen einer Abschlusssitzung zu besprechen.")
    
    add_separator(doc)
    add_field(doc, "Ort, Datum", 50)
    add_signature_line(doc, "Unterschrift Partner/in 1")
    add_signature_line(doc, "Unterschrift Partner/in 2")
    add_signature_line(doc, "Charlotte von Sichart, Therapeutin")
    
    add_footer(doc)
    path = os.path.join(os.path.dirname(__file__), 'Therapievertrag_Paartherapie_Charlotte_von_Sichart.docx')
    doc.save(path)
    return path

if __name__ == '__main__':
    p1 = create_schweigepflicht()
    print(f"Created: {p1}")
    p2 = create_rechnung()
    print(f"Created: {p2}")
    p3 = create_paartherapie()
    print(f"Created: {p3}")
