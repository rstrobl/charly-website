from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import subprocess

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

GREEN = RGBColor(0x2D, 0x4A, 0x3E)
GOLD = RGBColor(0xC9, 0xA8, 0x6C)

# Logo
logo_para = doc.add_paragraph()
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_para.add_run().add_picture('logo.jpg', width=Cm(4))
logo_para.paragraph_format.space_after = Pt(4)

# Gold separator
sep = doc.add_paragraph()
sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sep.add_run('━' * 50)
run.font.color.rgb = GOLD
run.font.size = Pt(8)
sep.paragraph_format.space_after = Pt(12)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Abrechnungsübersicht\nKostenerstattungsverfahren (§ 13 Abs. 3 SGB V)')
run.font.size = Pt(14)
run.font.color.rgb = GREEN
run.bold = True
title.paragraph_format.space_after = Pt(4)

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Charlotte von Sichart — Systemische Paartherapeutin')
run.font.size = Pt(10)
run.font.color.rgb = GOLD
sub.paragraph_format.space_after = Pt(6)

# Gold separator
sep2 = doc.add_paragraph()
sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sep2.add_run('━' * 50)
run.font.color.rgb = GOLD
run.font.size = Pt(8)
sep2.paragraph_format.space_after = Pt(12)

# Intro
intro = doc.add_paragraph()
run = intro.add_run('Alle Beträge bei 2,3-fachem Steigerungssatz (Regelhöchstsatz). ')
run.font.size = Pt(9)
run2 = intro.add_run('Eine Steigerung über den 2,3-fachen Satz wird von gesetzlichen Krankenkassen im Kostenerstattungsverfahren nicht übernommen.')
run2.font.size = Pt(9)
run2.bold = True
intro.paragraph_format.space_after = Pt(10)

def add_section_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN
    run.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table(rows_data):
    """rows_data: list of (ziffer, beschreibung, betrag) tuples"""
    table = doc.add_table(rows=len(rows_data)+1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    headers = ['GOÄ-Ziffer', 'Leistung', 'Betrag (2,3x)']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '2D4A3E'
        })
        shading.append(shading_elem)
    
    for row_idx, (ziffer, beschreibung, betrag) in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        row.cells[0].text = ziffer
        row.cells[1].text = beschreibung
        row.cells[2].text = betrag
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(9)
        row.cells[2].width = Cm(3)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_note(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.size = Pt(9)
        run.bold = True
    run = p.add_run(text)
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

def add_sum_line(text, amount):
    p = doc.add_paragraph()
    run = p.add_run(f'  ➤ {text}: ')
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = GREEN
    run2 = p.add_run(amount)
    run2.font.size = Pt(10)
    run2.bold = True
    p.paragraph_format.space_after = Pt(2)

# === KURZZEITTHERAPIE ===
add_section_heading('1. Kurzzeittherapie (bis 24 Sitzungen)')
add_table([
    ('812a', 'Psychotherapeutische Kurzzeittherapie\nEinzelbehandlung, mind. 50 Min.', '134,06 €'),
    ('801a', 'Erhebung des aktuellen psychischen Befundes\n(bei diagnostischer Begründung)', '33,52 €'),
])
add_sum_line('Maximale Summe pro Sitzung', '167,58 €')
add_note('Max. 48 x 25 Min. (= 24 x 50 Min.) pro Patient/Jahr', bold_prefix='Kontingent: ')
add_note('801a nur bei diagnostischer Begründung (Erstgespräch, Verlaufsdiagnostik), nicht jede Sitzung', bold_prefix='Hinweis: ')

# === LANGZEITTHERAPIE ===
add_section_heading('2. Langzeittherapie (ab Sitzung 25)')
add_table([
    ('870a', 'Psychotherapeutische Langzeittherapie\nEinzelbehandlung, mind. 50 Min.', '100,55 €'),
    ('801a', 'Erhebung des aktuellen psychischen Befundes\n(bei diagnostischer Begründung)', '33,52 €'),
])
add_sum_line('Maximale Summe pro Sitzung', '134,07 €')
add_note('Genehmigungspflichtig (Gutachterverfahren)', bold_prefix='Hinweis: ')

# === SPRECHSTUNDE ===
add_section_heading('3. Psychotherapeutische Sprechstunde')
add_table([
    ('812a', 'Psychotherapeutische Sprechstunde\nEinzelbehandlung, mind. 50 Min.', '134,06 €'),
])
add_sum_line('Summe pro Sitzung', '134,06 €')
add_note('Max. 6 x 25 Min. pro Patient/Jahr', bold_prefix='Kontingent: ')
add_note('801a NICHT kombinierbar mit Sprechstunde', bold_prefix='Achtung: ')

# === AKUTBEHANDLUNG ===
add_section_heading('4. Akutbehandlung')
add_table([
    ('812a', 'Psychotherapeutische Akutbehandlung\nEinzelbehandlung, mind. 50 Min.', '134,06 €'),
    ('801a', 'Erhebung des aktuellen psychischen Befundes', '33,52 €'),
])
add_sum_line('Maximale Summe pro Sitzung', '167,58 €')
add_note('Max. 24 x 25 Min. pro Patient/Jahr', bold_prefix='Kontingent: ')
add_note('Keine Vorabgenehmigung durch Krankenkasse nötig', bold_prefix='Vorteil: ')

# Gold separator
sep3 = doc.add_paragraph()
sep3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sep3.add_run('━' * 50)
run.font.color.rgb = GOLD
run.font.size = Pt(8)
sep3.paragraph_format.space_before = Pt(10)
sep3.paragraph_format.space_after = Pt(10)

# Important notes section
add_section_heading('Wichtige Hinweise')

notes = [
    ('⚠️ Kein 3,5x-Faktor: ', 'Im Kostenerstattungsverfahren darf maximal der 2,3-fache Steigerungssatz angesetzt werden. Höhere Sätze werden von GKV nicht erstattet.'),
    ('✅ Verfahrensübergreifend: ', 'Die Analogziffern 812a, 870a und 801a gelten für alle Richtlinienverfahren (TfP, VT, Systemische Therapie, Analytische PT).'),
    ('📅 Gültig seit: ', '01.07.2024 — Gemeinsame Abrechnungsempfehlungen von BÄK, BPtK, PKV und Beihilfe.'),
    ('💡 Empfehlung: ', 'Bei Kostenerstattung die neuen Analogziffern (812a/870a + 801a) bevorzugen — sie ergeben bei 2,3x oft mehr als die alten Ziffern (861/863) bei 3,5x.'),
]

for bold_part, normal_part in notes:
    p = doc.add_paragraph()
    run = p.add_run(bold_part)
    run.font.size = Pt(9)
    run.bold = True
    run = p.add_run(normal_part)
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)

# Comparison box
sep4 = doc.add_paragraph()
sep4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sep4.add_run('━' * 50)
run.font.color.rgb = GOLD
run.font.size = Pt(8)
sep4.paragraph_format.space_before = Pt(8)
sep4.paragraph_format.space_after = Pt(8)

add_section_heading('Vergleich: Alte vs. neue Ziffern (Einzeltherapie TfP)')
add_table([
    ('861 (alt)', 'TfP-Einzelbehandlung, 3,5x Faktor', '140,76 €'),
    ('812a (neu)', 'Kurzzeittherapie, 2,3x Faktor', '134,06 €'),
    ('812a + 801a', 'Kurzzeittherapie + Befund, 2,3x', '167,58 €'),
    ('870a (neu)', 'Langzeittherapie, 2,3x Faktor', '100,55 €'),
    ('870a + 801a', 'Langzeittherapie + Befund, 2,3x', '134,07 €'),
])

p = doc.add_paragraph()
run = p.add_run('➤ 812a + 801a (167,58 €) übertrifft die alte 861 bei 3,5x (140,76 €) deutlich!')
run.font.size = Pt(9)
run.bold = True
run.font.color.rgb = GREEN
p.paragraph_format.space_after = Pt(12)

# Footer
sep5 = doc.add_paragraph()
sep5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sep5.add_run('━' * 50)
run.font.color.rgb = GOLD
run.font.size = Pt(8)

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Charlotte von Sichart · Systemische Paartherapeutin\nKameruner Straße 43 · 13351 Berlin\nwww.von-sichart.de')
run.font.size = Pt(8)
run.font.color.rgb = GOLD

docx_path = 'Kostenerstattung_Abrechnungsuebersicht_Charlotte_von_Sichart.docx'
pdf_path = 'Kostenerstattung_Abrechnungsuebersicht_Charlotte_von_Sichart.pdf'

doc.save(docx_path)
print(f'DOCX saved: {docx_path}')

# Convert to PDF
subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', docx_path], check=True)
print(f'PDF saved: {pdf_path}')
